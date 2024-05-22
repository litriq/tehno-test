from django.shortcuts import get_object_or_404
from rest_framework import viewsets, status
from rest_framework.response import Response
from testing.models import Test, CompletedTest
from testing.serializers import TestListSerializer, TestSerializer, CompletedTestSerializer, UserSerializer
from django.views.generic import View
from django.http import HttpResponse, HttpResponseBadRequest, HttpResponseForbidden
from rest_framework.permissions import IsAuthenticated
from rest_framework.decorators import permission_classes
from rest_framework.views import APIView
from django.db.models import Exists, OuterRef
import csv
import docx
from django.http import HttpResponse
from datetime import datetime
from docx.shared import RGBColor
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from django.db.models import F


def fill_cells(cells, color_hex):
    color = RGBColor.from_string(color_hex)

    for cell in cells:
        tblCell = cell._tc
        tblCellProperties = tblCell.get_or_add_tcPr()
        clShading = OxmlElement('w:shd')
        # Hex of Dark Blue Shade {R:0x00, G:0x51, B:0x9E}
        clShading.set(qn('w:fill'), color_hex)
        tblCellProperties.append(clShading)


class TestViewSet(viewsets.ReadOnlyModelViewSet):
    permission_classes = [IsAuthenticated]

    queryset = Test.objects.all()
    serializer_class = TestListSerializer

    def list(self, request):
        user = request.user
        queryset = self.get_queryset()
        queryset = queryset.annotate(
            completed=Exists(
                CompletedTest.objects.filter(
                    user=user,
                    test=OuterRef('pk')
                )
            )
        )
        serializer = TestListSerializer(queryset, many=True)
        return Response(serializer.data)

    def retrieve(self, request, pk=None):
        test = get_object_or_404(Test, pk=pk)
        serializer = TestSerializer(test)
        return Response(serializer.data)


class CompleteTestView(APIView):
    permission_classes = [IsAuthenticated]

    def post(self, request, pk):
        test = get_object_or_404(Test, id=pk)
        user = request.user
        score = request.data.get('score')

        completed_test = CompletedTest.objects.create(
            user=user,
            test=test,
            score=score
        )

        return Response(CompletedTestSerializer(completed_test).data, status=201)


class UserView(APIView):
    permission_classes = [IsAuthenticated]

    def get(self, request):
        user = request.user

        roles = user.groups.all()
        user.full_name = user.get_full_name()

        serializer = UserSerializer(user)

        return Response(serializer.data)


class CompletedTestsDOCXView(View):
    permission_classes = [IsAuthenticated]

    def get(self, request):
        # получаем данные
        tests = CompletedTest.objects.order_by(F('score').desc())

        # создаем документ
        doc = docx.Document()

        # добавляем заголовок
        doc.add_heading('Отчет по пройденным тестам', 0)

        # добавляем таблицу
        table = doc.add_table(rows=1, cols=4)
        hdr_cells = table.rows[0].cells
        hdr_cells[0].text = 'Пользователь'
        hdr_cells[1].text = 'Тест'
        hdr_cells[2].text = 'Результат'
        hdr_cells[3].text = 'Дата'

        # заполняем строки таблицы данными
        for test in tests:

            date = datetime.strptime(
                str(test.completed_at), "%Y-%m-%d %H:%M:%S.%f%z")

            row_cells = table.add_row().cells
            row_cells[0].text = str(test.user.get_full_name())
            row_cells[1].text = str(test.test)
            row_cells[2].text = str(test.score) + "%"
            row_cells[3].text = date.strftime("%Y-%m-%d")

            if (test.score >= 80):
                fill_cells([row_cells[2]], "77E27C")

            if (test.score <= 40):
                fill_cells([row_cells[2]], "FF5858")

        # сохраняем документ и возвращаем его
        response = HttpResponse(
            content_type='application/vnd.openxmlformats-officedocument.wordprocessingml.document')
        response['Content-Disposition'] = 'attachment; filename=report.docx'
        doc.save(response)

        return response
